"""Renderers for the credit assessment report.

All three renderers consume the same block model produced by `builder.assemble_report`
so the branded output is identical across PDF, Word and HTML preview.
"""

from __future__ import annotations

from html import escape

from .builder import BRANDING

INK = "#16130F"
AMBER = "#CC8800"
PAPER = "#FBF9F4"


# --------------------------------------------------------------------------- #
# PDF (reportlab)
# --------------------------------------------------------------------------- #
def export_pdf(report: dict) -> bytes:
    from io import BytesIO

    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        ListFlowable,
        ListItem,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    cover = report["cover"]
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm, title="Credit Assessment Report",
    )

    ss = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=ss["Heading1"], textColor=colors.HexColor(INK), fontSize=16, spaceAfter=8)
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], textColor=colors.HexColor(AMBER), fontSize=13, spaceBefore=10, spaceAfter=6)
    h3 = ParagraphStyle("h3", parent=ss["Heading3"], textColor=colors.HexColor(INK), fontSize=11, spaceBefore=6, spaceAfter=3)
    body = ParagraphStyle("body", parent=ss["BodyText"], fontSize=9.5, leading=14, alignment=TA_LEFT)
    small = ParagraphStyle("small", parent=body, fontSize=8.5, textColor=colors.HexColor("#534B40"))

    def footer(canvas, d):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor(AMBER))
        canvas.setLineWidth(0.6)
        canvas.line(2 * cm, 1.4 * cm, A4[0] - 2 * cm, 1.4 * cm)
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor("#534B40"))
        canvas.drawString(2 * cm, 1.0 * cm, f"{BRANDING['name']}  ·  {BRANDING['title']}")
        canvas.drawRightString(A4[0] - 2 * cm, 1.0 * cm, BRANDING["email"])
        canvas.restoreState()

    flow = []
    # Cover
    flow.append(Spacer(1, 1.5 * cm))
    flow.append(Paragraph("CREDIT ASSESSMENT REPORT", ParagraphStyle("cv", parent=h1, fontSize=22, textColor=colors.HexColor(AMBER))))
    flow.append(Spacer(1, 0.4 * cm))
    flow.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor(AMBER)))
    flow.append(Spacer(1, 0.6 * cm))
    cover_rows = [
        ["Analyst", cover["analyst_name"]],
        ["Client", cover["company_name"]],
        ["Engagement question", cover["purpose"] or "—"],
        ["Internal rating", cover["rating_band"]],
        ["Date", cover["generated_at"]],
    ]
    ct = Table([[Paragraph(f"<b>{k}</b>", body), Paragraph(escape(str(v)), body)] for k, v in cover_rows],
               colWidths=[5 * cm, 11 * cm])
    ct.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LINEBELOW", (0, 0), (-1, -1), 0.3, colors.HexColor("#E3DDD2")),
    ]))
    flow.append(ct)
    flow.append(Spacer(1, 0.6 * cm))

    for b in report["blocks"]:
        k = b["kind"]
        if k == "pagebreak":
            flow.append(PageBreak())
        elif k == "h1":
            flow.append(Paragraph(escape(b["text"]), h1))
        elif k == "h2":
            flow.append(Paragraph(escape(b["text"]), h2))
        elif k == "h3":
            flow.append(Paragraph(escape(b["text"]), h3))
        elif k == "p":
            flow.append(Paragraph(escape(b["text"]), body))
            flow.append(Spacer(1, 3))
        elif k == "bullets":
            items = [ListItem(Paragraph(escape(it), body)) for it in b["items"]]
            flow.append(ListFlowable(items, bulletType="bullet", start="•", leftIndent=12))
            flow.append(Spacer(1, 3))
        elif k == "banner":
            bg = {"red": "#B23A3A", "amber": "#C18A1E", "green": "#3F7D4E"}.get(b.get("level", "amber"), "#C18A1E")
            bt = Table([[Paragraph(f"<b>{escape(b['text'])}</b>",
                                   ParagraphStyle("bn", parent=body, textColor=colors.white, fontSize=11)),
                         Paragraph(escape(b.get("detail", "")),
                                   ParagraphStyle("bnd", parent=small, textColor=colors.white))]],
                        colWidths=[8 * cm, 8 * cm])
            bt.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(bg)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ]))
            flow.append(bt)
            flow.append(Spacer(1, 6))
        elif k == "table":
            data = [[Paragraph(f"<b>{escape(h)}</b>", small) for h in b["headers"]]]
            for row in b["rows"]:
                data.append([Paragraph(escape(str(c)), small) for c in row])
            t = Table(data, colWidths=[4.5 * cm] + [3 * cm] * (len(b["headers"]) - 1))
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#16130F")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E3DDD2")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(PAPER)]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            flow.append(t)
            flow.append(Spacer(1, 6))

    doc.build(flow, onFirstPage=footer, onLaterPages=footer)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Word (python-docx)
# --------------------------------------------------------------------------- #
def export_docx(report: dict) -> bytes:
    from io import BytesIO

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    cover = report["cover"]
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    ink = RGBColor(0x16, 0x13, 0x0F)
    amber = RGBColor(0xCC, 0x88, 0x00)

    t = doc.add_paragraph("CREDIT ASSESSMENT REPORT")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(22)
    t.runs[0].font.color.rgb = amber

    for k, v in [("Analyst", cover["analyst_name"]), ("Client", cover["company_name"]),
                 ("Engagement question", cover["purpose"] or "—"),
                 ("Internal rating", cover["rating_band"]), ("Date", cover["generated_at"])]:
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(str(v))

    doc.add_paragraph("")

    for b in report["blocks"]:
        k = b["kind"]
        if k == "pagebreak":
            doc.add_page_break()
        elif k in ("h1", "h2", "h3"):
            p = doc.add_heading(b["text"], level=2 if k == "h2" else 3 if k == "h3" else 1)
            if k == "h2":
                p.runs[0].font.color.rgb = amber
        elif k == "p":
            doc.add_paragraph(b["text"])
        elif k == "bullets":
            for it in b["items"]:
                doc.add_paragraph(it, style="List Bullet")
        elif k == "banner":
            color = {"red": RGBColor(0xB2, 0x3A, 0x3A),
                     "amber": RGBColor(0xC1, 0x8A, 0x1E),
                     "green": RGBColor(0x3F, 0x7D, 0x4E)}.get(b.get("level", "amber"), RGBColor(0xC1, 0x8A, 0x1E))
            pb = doc.add_paragraph()
            run = pb.add_run(b["text"])
            run.bold = True
            run.font.color.rgb = color
            run.font.size = Pt(12)
            if b.get("detail"):
                pd = doc.add_paragraph()
                rd = pd.add_run(b["detail"])
                rd.font.color.rgb = color
                rd.font.size = Pt(9)
        elif k == "table":
            table = doc.add_table(rows=1, cols=len(b["headers"]))
            table.style = "Light Grid Accent 1"
            for i, h in enumerate(b["headers"]):
                table.rows[0].cells[i].text = h
            for row in b["rows"]:
                cells = table.add_row().cells
                for i, c in enumerate(row):
                    cells[i].text = str(c)
            doc.add_paragraph("")

    # footer branding
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = f"{BRANDING['name']}  ·  {BRANDING['title']}  ·  {BRANDING['email']}"

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# HTML preview
# --------------------------------------------------------------------------- #
def export_html(report: dict) -> str:
    cover = report["cover"]
    out = [f"<h1>Credit Assessment Report</h1>",
           f"<p><b>Analyst:</b> {escape(cover['analyst_name'])}<br>"
           f"<b>Client:</b> {escape(cover['company_name'])}<br>"
           f"<b>Engagement question:</b> {escape(cover['purpose'] or '—')}<br>"
           f"<b>Internal rating:</b> {escape(cover['rating_band'])}<br>"
           f"<b>Date:</b> {escape(cover['generated_at'])}</p><hr>"]
    for b in report["blocks"]:
        k = b["kind"]
        if k == "pagebreak":
            out.append("<hr style='page-break-before:always;border:none;border-top:1px dashed #CC8800;margin:24px 0'>")
        elif k == "h1":
            out.append(f"<h1>{escape(b['text'])}</h1>")
        elif k == "h2":
            out.append(f"<h2>{escape(b['text'])}</h2>")
        elif k == "h3":
            out.append(f"<h3>{escape(b['text'])}</h3>")
        elif k == "p":
            out.append(f"<p>{escape(b['text'])}</p>")
        elif k == "bullets":
            out.append("<ul>" + "".join(f"<li>{escape(i)}</li>" for i in b["items"]) + "</ul>")
        elif k == "banner":
            color = {"red": "#B23A3A", "amber": "#C18A1E", "green": "#3F7D4E"}.get(b.get("level", "amber"), "#C18A1E")
            out.append(
                f"<div style='background:{color};color:#fff;padding:12px 16px;border-radius:8px;margin:12px 0'>"
                f"<strong>{escape(b['text'])}</strong><br><span style='font-size:12px'>{escape(b.get('detail', ''))}</span></div>")
        elif k == "table":
            head = "".join(f"<th>{escape(h)}</th>" for h in b["headers"])
            rows = "".join("<tr>" + "".join(f"<td>{escape(str(c))}</td>" for c in r) + "</tr>"
                           for r in b["rows"])
            out.append(f"<table border='1' cellspacing='0' cellpadding='4'><thead><tr>{head}</tr></thead><tbody>{rows}</tbody></table>")
    out.append(f"<hr><p><small>{escape(BRANDING['name'])} · {escape(BRANDING['title'])} · {escape(BRANDING['email'])}</small></p>")
    return "\n".join(out)
